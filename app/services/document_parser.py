"""Document parser service for extracting text from various file formats.

Supports:
- Word documents (.docx)
- PDF files (.pdf)
- Text files (.txt)
- Markdown files (.md)
- Excel/CSV files (.xlsx, .csv)
"""
import logging
import io
import csv
from typing import Optional, Dict, List
from pathlib import Path
import re

log = logging.getLogger(__name__)


def parse_document(file_content: bytes, filename: str) -> Optional[str]:
    """
    Parse document and extract text content.
    
    Args:
        file_content: Raw bytes of the uploaded file
        filename: Original filename with extension
    
    Returns:
        Extracted text content or None if parsing fails
    """
    try:
        # Get file extension
        extension = Path(filename).suffix.lower()
        
        log.info(f"Parsing document: {filename} (extension: {extension})")
        
        # Route to appropriate parser based on extension
        if extension == '.docx':
            return parse_docx(file_content)
        elif extension == '.pdf':
            return parse_pdf(file_content)
        elif extension in ['.txt', '.text']:
            return parse_text(file_content)
        elif extension in ['.md', '.markdown']:
            return parse_markdown(file_content)
        elif extension in ['.csv', '.xlsx', '.xls']:
            return parse_spreadsheet(file_content, extension)
        else:
            log.warning(f"Unsupported file format: {extension}")
            return None
    
    except Exception as e:
        log.error(f"Error parsing document {filename}: {e}", exc_info=True)
        return None


def parse_docx(file_content: bytes) -> Optional[str]:
    """Parse Word document (.docx) and extract text."""
    try:
        from docx import Document
        
        # Load document from bytes
        doc = Document(io.BytesIO(file_content))
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        result = "\n\n".join(text_parts)
        log.info(f"✅ Extracted {len(result)} characters from DOCX")
        return result if result.strip() else None
    
    except ImportError:
        log.error("python-docx not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        log.error(f"Error parsing DOCX: {e}")
        return None


def parse_pdf(file_content: bytes) -> Optional[str]:
    """Parse PDF and extract text."""
    try:
        import pdfplumber
        
        # Try pdfplumber first (better quality)
        try:
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text.strip())
                    
                    # Also extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row and any(cell for cell in row if cell):
                                row_text = " | ".join(str(cell) for cell in row if cell)
                                text_parts.append(row_text)
            
            result = "\n\n".join(text_parts)
            log.info(f"✅ Extracted {len(result)} characters from PDF (pdfplumber)")
            return result if result.strip() else None
        
        except Exception as pdf_error:
            log.warning(f"pdfplumber failed, trying PyPDF2: {pdf_error}")
    
    except ImportError:
        log.warning("pdfplumber not installed, trying PyPDF2")
    
    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())
        
        result = "\n\n".join(text_parts)
        log.info(f"✅ Extracted {len(result)} characters from PDF (PyPDF2)")
        return result if result.strip() else None
    
    except ImportError:
        log.error("PyPDF2 not installed. Run: pip install PyPDF2")
        return None
    except Exception as e:
        log.error(f"Error parsing PDF: {e}")
        return None


def parse_text(file_content: bytes) -> Optional[str]:
    """Parse plain text file."""
    try:
        # Try UTF-8 first
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            text = file_content.decode('latin-1')
        
        log.info(f"✅ Extracted {len(text)} characters from TXT")
        return text.strip() if text.strip() else None
    
    except Exception as e:
        log.error(f"Error parsing text file: {e}")
        return None


def parse_markdown(file_content: bytes) -> Optional[str]:
    """Parse Markdown file and extract text."""
    try:
        # Try UTF-8 first
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')
        
        # Remove markdown syntax for cleaner extraction
        try:
            import markdown
            from html import unescape
            import re
            
            # Convert markdown to HTML then strip tags
            html = markdown.markdown(text)
            # Remove HTML tags
            text_only = re.sub(r'<[^>]+>', '', html)
            # Unescape HTML entities
            text_only = unescape(text_only)
            
            log.info(f"✅ Extracted {len(text_only)} characters from Markdown")
            return text_only.strip() if text_only.strip() else None
        
        except ImportError:
            # Fallback: just remove basic markdown syntax
            text = text.replace('#', '').replace('*', '').replace('_', '')
            log.info(f"✅ Extracted {len(text)} characters from Markdown (basic)")
            return text.strip() if text.strip() else None
    
    except Exception as e:
        log.error(f"Error parsing markdown: {e}")
        return None


def parse_spreadsheet(file_content: bytes, extension: str) -> Optional[str]:
    """Parse Excel/CSV and extract text."""
    try:
        if extension == '.csv':
            # Parse CSV
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')
            
            reader = csv.reader(io.StringIO(text))
            rows = []
            for row in reader:
                if row and any(cell.strip() for cell in row if cell):
                    rows.append(" | ".join(cell.strip() for cell in row if cell))
            
            result = "\n".join(rows)
            log.info(f"✅ Extracted {len(result)} characters from CSV")
            return result if result.strip() else None
        
        else:
            # Parse Excel (.xlsx, .xls)
            try:
                from openpyxl import load_workbook
                
                workbook = load_workbook(io.BytesIO(file_content), read_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"Sheet: {sheet_name}")
                    
                    for row in sheet.iter_rows(values_only=True):
                        if row and any(cell for cell in row if cell):
                            row_text = " | ".join(str(cell) for cell in row if cell)
                            text_parts.append(row_text)
                    
                    text_parts.append("")  # Blank line between sheets
                
                result = "\n".join(text_parts)
                log.info(f"✅ Extracted {len(result)} characters from Excel")
                return result if result.strip() else None
            
            except ImportError:
                log.error("openpyxl not installed. Run: pip install openpyxl")
                return None
    
    except Exception as e:
        log.error(f"Error parsing spreadsheet: {e}")
        return None


def extract_qa_with_gpt(
    document_text: str,
    openai_api_key: str = "",
    gemini_api_key: str = "",
    model_override: Optional[str] = None,
) -> List[Dict[str, any]]:
    """
    Use configured LLM to intelligently extract Q&A pairs from document text.
    
    Args:
        document_text: Raw text extracted from document
        openai_api_key: OpenAI API key
        gemini_api_key: Gemini API key (uses OpenAI-compatible endpoint)
        model_override: Optional model name override
    
    Returns:
        List of dictionaries with question, answer, and keywords
    """
    try:
        from openai import OpenAI

        openai_api_key = (openai_api_key or "").strip()
        gemini_api_key = (gemini_api_key or "").strip()

        if gemini_api_key:
            client = OpenAI(
                api_key=gemini_api_key,
                base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
            )
            model_name = model_override or "gemini-2.0-flash"
            provider = "gemini"
        elif openai_api_key:
            client = OpenAI(api_key=openai_api_key)
            model_name = model_override or "gpt-4o"
            provider = "openai"
        else:
            log.warning("No AI API key configured for document extraction")
            return []
        
        # Truncate if too long (GPT-4o has 128k context, but we'll be conservative)
        max_chars = 50000  # ~12,500 tokens
        if len(document_text) > max_chars:
            log.warning(f"Document too long ({len(document_text)} chars), truncating to {max_chars}")
            document_text = document_text[:max_chars] + "\n\n[Document truncated...]"
        
        system_prompt = """You are an expert at analyzing documents and extracting knowledge into Q&A format.

Your task:
1. Read the document carefully
2. Identify key information that customers might ask about
3. Create clear, specific Q&A pairs
4. Extract relevant keywords for each Q&A

Guidelines:
- Create 5-15 Q&A pairs (depending on document length)
- Questions should be natural (how customers would ask)
- Answers should be concise but complete (2-4 sentences)
- Include important details, numbers, dates, policies
- Keywords should help match customer queries

Return ONLY valid JSON in this exact format (no markdown, no extra text):
[
  {
    "question": "What are your business hours?",
    "answer": "We are open Monday to Friday from 9 AM to 6 PM EST, and Saturday from 10 AM to 4 PM EST. We are closed on Sundays and major holidays.",
    "keywords": ["hours", "open", "time", "schedule", "when", "available"]
  }
]"""

        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract Q&A pairs from this document:\n\n{document_text}"}
            ],
            temperature=0.3,  # Lower temperature for consistent extraction
            max_tokens=4000,
            timeout=60.0
        )
        
        result_text = response.choices[0].message.content or ""

        def _normalize_qa_pairs(raw_pairs: object) -> List[Dict[str, any]]:
            if not isinstance(raw_pairs, list):
                return []

            normalized: List[Dict[str, any]] = []
            for item in raw_pairs:
                if not isinstance(item, dict):
                    continue

                question = str(item.get("question", "")).strip()
                answer = str(item.get("answer", "")).strip()
                if not question or not answer:
                    continue

                keywords_raw = item.get("keywords", [])
                if isinstance(keywords_raw, list):
                    keywords = [str(k).strip() for k in keywords_raw if str(k).strip()]
                elif isinstance(keywords_raw, str):
                    keywords = [k.strip() for k in keywords_raw.split(",") if k.strip()]
                else:
                    keywords = []

                normalized.append({
                    "question": question,
                    "answer": answer,
                    "keywords": keywords,
                })

            return normalized

        # Parse JSON response with multiple fallbacks for model formatting variations.
        import json
        import re

        parse_candidates = [result_text.strip()]

        # Markdown fenced block candidate.
        code_block_match = re.search(r"```(?:json)?\s*(.*?)\s*```", result_text, re.DOTALL)
        if code_block_match:
            parse_candidates.append(code_block_match.group(1).strip())

        # First JSON array candidate anywhere in text.
        first_bracket = result_text.find("[")
        last_bracket = result_text.rfind("]")
        if first_bracket != -1 and last_bracket != -1 and last_bracket > first_bracket:
            parse_candidates.append(result_text[first_bracket:last_bracket + 1].strip())

        for candidate in parse_candidates:
            if not candidate:
                continue
            try:
                parsed = json.loads(candidate)
                qa_pairs = _normalize_qa_pairs(parsed)
                if qa_pairs:
                    log.info(f"✅ AI extracted {len(qa_pairs)} Q&A pairs via {provider}")
                    return qa_pairs
            except json.JSONDecodeError:
                continue

        log.error("Failed to parse AI extraction response as valid Q&A JSON")
        return []
    
    except Exception as e:
        log.error(f"Error extracting Q&A with AI: {e}", exc_info=True)
        return []


def extract_qa_fallback(document_text: str, max_pairs: int = 12) -> List[Dict[str, any]]:
    """
    Build Q&A pairs from plain text when AI extraction fails.

    Heuristic strategy:
    - Split into meaningful paragraphs
    - Turn each paragraph into a likely customer question + concise answer
    - Generate basic keywords from paragraph terms
    """
    if not document_text or not isinstance(document_text, str):
        return []

    # Normalize whitespace and split by paragraph breaks.
    cleaned = re.sub(r"\r\n?", "\n", document_text).strip()
    blocks = [b.strip() for b in re.split(r"\n\s*\n+", cleaned) if b.strip()]
    if not blocks:
        return []

    stopwords = {
        "the", "and", "for", "with", "that", "this", "from", "your", "you", "are",
        "our", "can", "will", "into", "about", "have", "has", "was", "were", "how",
        "what", "when", "where", "which", "their", "them", "more", "than", "just",
    }

    qa_pairs: List[Dict[str, any]] = []
    for block in blocks:
        if len(qa_pairs) >= max_pairs:
            break

        # Skip tiny fragments.
        if len(block) < 35:
            continue

        # Use first sentence as question seed.
        sentences = re.split(r"(?<=[.!?])\s+", block)
        first_sentence = (sentences[0] or "").strip()
        if not first_sentence:
            continue

        # If sentence is already a question, keep it. Otherwise convert to FAQ style.
        if first_sentence.endswith("?"):
            question = first_sentence
        else:
            stem = re.sub(r"[^\w\s-]", "", first_sentence).strip()
            if not stem:
                continue
            question = f"What should I know about {stem[:90]}?"

        answer = block[:450].strip()
        if len(answer) < 30:
            continue

        words = re.findall(r"[A-Za-z][A-Za-z0-9_-]{2,}", block.lower())
        unique_words: List[str] = []
        for w in words:
            if w in stopwords or w in unique_words:
                continue
            unique_words.append(w)
            if len(unique_words) >= 6:
                break

        qa_pairs.append({
            "question": question,
            "answer": answer,
            "keywords": unique_words,
        })

    return qa_pairs
